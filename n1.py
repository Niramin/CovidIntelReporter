import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import docx
from datetime import date



def main():
    sheet_url="https://docs.google.com/spreadsheets/d/1FyZft52WAj2qqNuZTq-WrRrJy0BRs2nEOjy_CQGqvgk/edit#gid=0"
    url_1 = sheet_url.replace("/edit#gid=" , "/export?format=csv&gid=")
    df=pd.read_csv(url_1)
    #print(df.iloc[:,4].values)




    #Fever Duration
    fev_duration=df.iloc[:,4].values
    fev_duration_mean=np.mean(fev_duration)
    fev_duration_iqr=np.percentile(fev_duration,75)-np.percentile(fev_duration,25)
    fev_duration_median=np.median(fev_duration)

    #Fever Temperature
    fev_temp=df.iloc[:,5].values
    fev_temp_iqr=np.percentile(fev_temp,75)-np.percentile(fev_temp,25)
    fev_temp_mean=np.mean(fev_temp)
    fev_temp_median=np.median(fev_temp)

    #Age
    age=df.iloc[:,6].values
    age_iqr=np.percentile(age,75)-np.percentile(age,25)
    age_mean=np.mean(age)
    age_median=np.median(age)

    #BMI
    bmi=df.iloc[:,7].values
    bmi_iqr=np.percentile(bmi,75)-np.percentile(bmi,25)
    bmi_mean=np.mean(bmi)
    bmi_median=np.median(bmi)


    #Oxygen Saturation Level
    osat=df.iloc[:,10]
    osat_mean=np.mean(osat)
    osat_median=np.median(osat)
    osat_iqr=np.percentile(osat,75)-np.percentile(osat,25)

    #Diabetic fraction
    dia=df.iloc[:,8].values
    tl=dia.size
    dia[dia=="Yes"]=25
    yed=np.count_nonzero(dia==25)
    yed=yed/tl

    #Heart_Problems fraction
    hp=df.iloc[:,9].values
    tlh=hp.size
    hp[hp=="Yes"]=25
    yeh=np.count_nonzero(hp==25)
    yeh=yeh/tlh

    #Asthma
    ast=df.iloc[:,11].values
    astl=ast.size
    ast[ast=="Yes"]=25
    yea=np.count_nonzero(ast==25)
    yea=yea/astl

    #Cough
    cou=df.iloc[:,12].values
    coul=cou.size
    cou[cou=="Yes"]=25
    yec=np.count_nonzero(cou==25)
    yec=yec/coul

    #ThroatAche
    tha=df.iloc[:,13].values
    thal=tha.size
    tha[tha=="Yes"]=25
    yet=np.count_nonzero(tha==25)
    yet=yet/thal

    #City Count
    cit=df.iloc[:,2].values
    d=dict()
    for i in cit:
        if i in d.keys():
            d[i]+=1
        else:
            d[i]=1
    #print(d.keys(),d.values())
    xd=[x+"\n" for x in d.keys()]
    plt.bar(d.keys(),d.values())
    plt.xticks(rotation=90)
    plt.savefig("outputcity.jpg",bbox_inches='tight')
    plt.clf()
    sorted_city=max(zip(d.values(), d.keys()))
    

    #State Count
    sta=df.iloc[:,3].values
    sd=dict()
    for i in sta:
        if i in sd.keys():
            sd[i]+=1
        else:
            sd[i]=1
    plt.bar(sd.keys(),sd.values())
    plt.xticks(rotation=90)
    plt.savefig("outputstate.jpg",bbox_inches='tight')
    plt.clf()
    sorted_state=max(zip(sd.values(), sd.keys()))


    f=open("Parameters.csv","w")
    #Saving the damn report
    doc=docx.Document()
    tdate=str(date.today())
    doc.add_heading('Health Intel Report '+tdate, 0)
    doc.add_heading('Individual Report', level=1)
    #Numeric
    doc.add_paragraph('Numeric Data Analysis', style='Intense Quote')
    
    #Fever Duration
    doc.add_heading("Fever Duration",level=2)
    doc.add_paragraph(

    'The average fever duration is: '+"%.2f"%(fev_duration_mean)+" days",style='List Bullet'

    )

    doc.add_paragraph(

    'The median fever duration is: '+"%.2f"%(fev_duration_median)+" days",style='List Bullet'

    )

    doc.add_paragraph(

    'The interquartile range for fever duration: '+"%.2f"%(fev_duration_iqr)+" days",style='List Bullet'

    )
    s=""
    s+=("%.2f"%(fev_duration_median)+","+"%.2f"%(fev_duration_iqr)+",")

    #Fever Temperature
    doc.add_heading("Fever Temperature",level=2)

    doc.add_paragraph(

    'The average fever temperature reported is: '+"%.2f"%(fev_temp_mean)+" \u2103 F",style='List Bullet'

    )

    doc.add_paragraph(

    'The median fever temperature is: '+"%.2f"%(fev_temp_median)+"\u2103 F",style='List Bullet'

    )

    doc.add_paragraph(

    'The interquartile range for fever temperature: '+"%.2f"%(fev_temp_iqr)+"\u2103 F",style='List Bullet'

    )
    s+=("%.2f"%(fev_temp_median)+","+"%.2f"%(fev_temp_iqr)+",")

    #Age
    doc.add_heading("Age",level=2)

    doc.add_paragraph(

    'The average age of patient is : '+"%.2f"%(age_mean)+" years",style='List Bullet'

    )

    doc.add_paragraph(

    'The median age of pateint is : '+"%.2f"%(age_median)+" years",style='List Bullet'

    )

    doc.add_paragraph(

    'The interquartile range for age is : '+"%.2f"%(age_iqr)+" years",style='List Bullet'

    )
    s="%.2f"%(age_median)+","+"%.2f"%(age_iqr)+","

    #BMI
    doc.add_heading("BMI (Body Mass Undex)",level=2)

    doc.add_paragraph(

    'The average BMI of patient is : '+"%.2f"%(bmi_mean),style='List Bullet'

    )
    


    doc.add_paragraph(

    'The median value is : '+"%.2f"%(bmi_median),style='List Bullet'

    )

    doc.add_paragraph(

    'The interquartile range for BMI is : '+"%.2f"%(bmi_iqr),style='List Bullet'

    )
    s+=("%.2f"%(bmi_median)+","+"%.2f"%(bmi_iqr)+",")
    

    #Oxygen Saturation Level
    doc.add_heading("Oxygen Saturation Level",level=2)

    doc.add_paragraph(

    'The average oxygen saturation level is : '+"%.2f"%(osat_mean)+"%",style='List Bullet'

    )

    doc.add_paragraph(

    'The median level is : '+"%.2f"%(osat_median)+"%",style='List Bullet'

    )

    doc.add_paragraph(

    'The interquartile range is : '+"%.2f"%(osat_iqr)+"% points",style='List Bullet'

    )
    s+=("%.2f"%(osat_mean)+",")

    #Nominal Data

    doc.add_paragraph('Nominal Data Analysis', style='Intense Quote')

    #Diabetic fraction
    doc.add_paragraph(

    "%.2f"%(yed*100)+"% Covid positive patients reported being Diabetic",style='List Bullet'

    )
    s+=("Diabetes,"+str(yed)+",")


    #Heart Problems Fraction
    doc.add_paragraph(

    "%.2f"%(yeh*100)+"% Covid positive patients reported having Heart Problems",style='List Bullet'

    )
    s+=(str(yeh)+",")
    

    #Asthma
    doc.add_paragraph(

    "%.2f"%(yea*100)+"% Covid positive patients reported being Asthmatic",style='List Bullet'

    )
    s+=(str(yea)+",")

    #Cough
    doc.add_paragraph(

    "%.2f"%(yec*100)+"% Covid positive patients reported having Cough",style='List Bullet'

    )
    s+=(str(yec)+",")
    

    #Throat Ache
    doc.add_paragraph(

    "%.2f"%(yet*100)+"% Covid positive patients reported having Throat Ache",style='List Bullet'

    )
    s+=(str(yet)+",")


    #City Report
    doc.add_heading('City Report', level=1)
    doc.add_picture(r'C:\Users\Shashwat Ratna\Desktop\realshh\Learn\outputcity.jpg')
    doc.add_paragraph('Outbreak Alert', style='Intense Quote')
    doc.add_paragraph(

    str(sorted_city[1]) + " is most vulnerable with "+str(d[sorted_city[1]])+" cases recently reported",style='List Bullet'

    )
    s+=(sorted_city[1]+",")
    #State Report
    doc.add_heading('State Report', level=1)
    doc.add_picture(r'C:\Users\Shashwat Ratna\Desktop\realshh\Learn\outputstate.jpg')
    doc.add_paragraph('Pandemic Alert', style='Intense Quote')
    doc.add_paragraph(

    str(sorted_state[1]) + " is most vulnerable with "+str(sd[sorted_state[1]])+" cases recently reported",style='List Bullet'

    )
    s+=(sorted_state[1]+",")

    paragraph = doc.add_paragraph("\n\nEnd of Report")
    paragraph.alignment = 1 # for left, 1 for center, 2 right, 3 justify ..

    paragraph2 = doc.add_paragraph("_________________________________")
    paragraph2.alignment = 1 # for left, 1 for center, 2 right, 3 justify ..

    f.writelines(s)



    

    
    f.close()

    doc.save("demo.docx")









if __name__=="__main__":
    main()

import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import docx
from datetime import date
import emailr


def makefile(tf,fd,ft,ag,bm,osatl,ci,st):
    doc=docx.Document()
    tdate=str(date.today())
    doc.add_heading('Individual Health Report '+tdate, 0)    
    paragraph2 = doc.add_paragraph("for "+ str(tf['Name'].iloc[0]))
    paragraph2.alignment = 1

    f=open(r"C:\Users\Shashwat Ratna\Desktop\realshh\Learn\Parameters.csv","r")
    sli=f.read().split(",")
    print(sli)
    
        


    doc.save("demoi.docx")

def main():
    #sheet_url="https://docs.google.com/spreadsheets/d/1TeEHPHd4lMdZpIYaSjLJuEv7CP6e0wdnyZJ1BFSexCY/edit#gid=1873632045"
    #url_1 = sheet_url.replace("/edit#gid=" , "/export?format=csv&gid=")
    #userdata=pd.read_csv(url_1)

    sheet_url="https://docs.google.com/spreadsheets/d/1tz-0vENVAirieoaa-ngU6A992lsmCMd4Kz0RyoS7icU/edit#gid=1195946014"
    url_1 = sheet_url.replace("/edit#gid=" , "/export?format=csv&gid=")
    userreq=pd.read_csv(url_1)

    cir=userreq
    #Handling 0 request, ie Covid Intel Request
    for i in range(len(cir)):
      emailr.emailer( list(cir.iloc[i])[2] ,"demo.docx",r"C:\Users\Shashwat Ratna\Desktop\realshh\Learn\demo.docx")
       

    
if __name__=="__main__":
    main()